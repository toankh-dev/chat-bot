"""
Document Processing Service - Extract text from various file formats.
"""

from typing import BinaryIO, Optional
import io
import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


class DocumentProcessingService:
    """Service for extracting text content from various document formats."""

    def __init__(self):
        self.supported_formats = {".pdf", ".docx", ".txt", ".md"}

    async def extract_text(self, file_content: BinaryIO, filename: str, content_type: str) -> str:
        """
        Extract text content from a document file.

        Args:
            file_content: Binary file content
            filename: Original filename
            content_type: MIME type of the file

        Returns:
            Extracted text content as string

        Raises:
            ValueError: If file format is not supported
        """
        file_extension = os.path.splitext(filename)[1].lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Reset file pointer to beginning
        file_content.seek(0)

        if file_extension == ".pdf":
            return await self._extract_from_pdf(file_content)
        elif file_extension == ".docx":
            return await self._extract_from_docx(file_content)
        elif file_extension in {".txt", ".md"}:
            return await self._extract_from_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    async def _extract_from_pdf(self, file_content: BinaryIO) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PdfReader(file_content)
            text_parts = []

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    print(f"Error extracting text from page {page_num + 1}: {e}")
                    continue

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise ValueError("No text content could be extracted from PDF")

            return full_text

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    async def _extract_from_docx(self, file_content: BinaryIO) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_content)
            text_parts = []

            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise ValueError("No text content could be extracted from DOCX")

            return full_text

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    async def _extract_from_text(self, file_content: BinaryIO) -> str:
        """Extract text from plain text files (TXT, MD)."""
        try:
            # Try UTF-8 first
            try:
                text = file_content.read().decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                file_content.seek(0)
                text = file_content.read().decode('latin-1')

            if not text.strip():
                raise ValueError("File is empty")

            return text

        except Exception as e:
            raise ValueError(f"Failed to extract text from file: {str(e)}")

    def validate_extracted_text(self, text: str, min_length: int = 10) -> bool:
        """
        Validate that extracted text is meaningful.

        Args:
            text: Extracted text
            min_length: Minimum text length

        Returns:
            True if text is valid, False otherwise
        """
        if not text or not text.strip():
            return False

        if len(text.strip()) < min_length:
            return False

        return True

    def get_text_statistics(self, text: str) -> dict:
        """
        Get statistics about extracted text.

        Args:
            text: Extracted text

        Returns:
            Dictionary with text statistics
        """
        words = text.split()
        lines = text.split('\n')

        return {
            "character_count": len(text),
            "word_count": len(words),
            "line_count": len(lines),
            "non_whitespace_count": len(text.replace(' ', '').replace('\n', '').replace('\t', ''))
        }
